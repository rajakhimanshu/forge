"""
Layer 5 — Docx Export Agent
=============================
Packages the full pipeline state into a formatted, downloadable Word document.
Structure mirrors the forge_build_guide_v2.docx gold standard:
  1. Cover page — project name, date, env profile, services checklist
  2. Pre-build checklist — all API keys, signups, installs
  3. Project setup — folder creation, .env template
  4. Feature-by-feature SETUP STEP / CODING STEP sequences
  5. Testing guide
  6. Deployment guide
  7. GTM launch checklist

Uses python-docx (pure Python, no Node.js required).
SETUP step badges: blue background. CODING step badges: green background.
"""
from tools.llm_router import safe_print
import os
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    safe_print("[DOCX EXPORTER] python-docx not installed — install with: pip install python-docx")

load_dotenv()

# ── Color constants ───────────────────────────────────────────────────────────
BLUE_BG   = RGBColor(0x1E, 0x6F, 0xB5) if DOCX_AVAILABLE else None   # SETUP badge
GREEN_BG  = RGBColor(0x1B, 0x7A, 0x40) if DOCX_AVAILABLE else None   # CODING badge
DARK_BG   = RGBColor(0x1A, 0x1A, 0x2E) if DOCX_AVAILABLE else None   # Cover
ACCENT    = RGBColor(0x00, 0xD4, 0x8A) if DOCX_AVAILABLE else None    # Accent
WHITE     = RGBColor(0xFF, 0xFF, 0xFF) if DOCX_AVAILABLE else None
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5) if DOCX_AVAILABLE else None
CODE_FONT = "Courier New"
BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"


def _set_cell_bg(cell, rgb_color):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb_color)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_step_badge(doc: 'Document', step_type: str, label: str, step_number: int):
    """Add a coloured SETUP/CODING step badge row."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Badge cell (SETUP=blue, CODING=green)
    badge_cell = tbl.rows[0].cells[0]
    color = BLUE_BG if step_type == "setup" else GREEN_BG
    _set_cell_bg(badge_cell, color)
    badge_run = badge_cell.paragraphs[0].add_run(
        f"{'SETUP' if step_type == 'setup' else 'CODING'} STEP {step_number}"
    )
    badge_run.bold = True
    badge_run.font.color.rgb = WHITE
    badge_run.font.size = Pt(10)
    badge_run.font.name = BODY_FONT

    # Label cell
    label_cell = tbl.rows[0].cells[1]
    _set_cell_bg(label_cell, LIGHT_GREY)
    label_run = label_cell.paragraphs[0].add_run(f"  {label}")
    label_run.bold = True
    label_run.font.size = Pt(11)
    label_run.font.name = BODY_FONT

    # Column widths
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(5.0)
    doc.add_paragraph()


def _add_code_block(doc: 'Document', code: str):
    """Add a monospace code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = CODE_FONT
    run.font.size = Pt(9)
    # grey shading for code block
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)


def _add_heading(doc: 'Document', text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = HEADING_FONT
    return p


def _build_service_table(doc: 'Document', service_bundle):
    """Add the pre-build services checklist table."""
    if not service_bundle or not service_bundle.services:
        doc.add_paragraph("No external services resolved for this project.")
        return

    _add_heading(doc, "Services Checklist", level=2)
    doc.add_paragraph(
        "Complete ALL signups and get ALL API keys BEFORE writing any code."
    ).italic = True

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Light Grid Accent 1'

    headers = ["Service", "Provider", "Free Tier", "Env Var", "Sign Up URL"]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    for svc in service_bundle.services:
        row = tbl.add_row()
        row.cells[0].text = svc.infra_need
        row.cells[1].text = svc.recommended_service
        row.cells[2].text = svc.current_free_tier[:60] if svc.current_free_tier else ""
        row.cells[3].text = svc.env_var_name
        row.cells[4].text = svc.signup_url

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()

    # Install commands
    _add_heading(doc, "Install Commands", level=3)
    for svc in service_bundle.services:
        if svc.sdk_install_cmd:
            doc.add_paragraph(f"# {svc.infra_need}")
            _add_code_block(doc, svc.sdk_install_cmd)
    doc.add_paragraph()


def _build_env_template(doc: 'Document', service_bundle):
    """Add .env template with all resolved env vars pre-filled."""
    _add_heading(doc, ".env File Template", level=2)
    doc.add_paragraph(
        "Create a file named .env in your project root and paste this content:"
    )
    if service_bundle and service_bundle.services:
        env_lines = ["# Copy this file to .env and fill in your API keys\n"]
        for svc in service_bundle.services:
            env_lines.append(f"# {svc.infra_need} — get from: {svc.api_key_location}")
            env_lines.append(f"{svc.env_var_name}=your_{svc.env_var_name.lower()}_here\n")
        _add_code_block(doc, "\n".join(env_lines))
    else:
        _add_code_block(doc, "# No external services configured\nPORT=3000\nNODE_ENV=development")
    doc.add_paragraph()


def _build_gtm_section(doc: 'Document', gtm_output, roadmap_output):
    """Add the GTM launch checklist section."""
    _add_heading(doc, "Go-To-Market Launch Checklist", level=1)

    if gtm_output and not getattr(gtm_output, 'failed', False):
        _add_heading(doc, "Primary Channel", level=2)
        doc.add_paragraph(gtm_output.primary_channel or "TBD")

        _add_heading(doc, "Cold Outreach Script", level=2)
        doc.add_paragraph(gtm_output.cold_outreach_script or "TBD")

        for week_num, week_actions in enumerate([
            gtm_output.week1_actions,
            gtm_output.week2_actions,
            gtm_output.week3_actions,
            gtm_output.week4_money_ask,
        ], 1):
            if week_actions:
                _add_heading(doc, f"Week {week_num} Actions", level=2)
                for action in week_actions:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(action)

        if gtm_output.viral_mechanic:
            _add_heading(doc, "Viral Mechanic", level=2)
            doc.add_paragraph(gtm_output.viral_mechanic)

    if roadmap_output and not getattr(roadmap_output, 'failed', False):
        _add_heading(doc, "30-Day Action Plan", level=2)
        if roadmap_output.days_1_7:
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = 'Light Grid'
            for h, cell in zip(["Day", "Action", "Platform"], tbl.rows[0].cells):
                cell.paragraphs[0].add_run(h).bold = True

            for da in roadmap_output.days_1_7:
                row = tbl.add_row()
                row.cells[0].text = str(da.day)
                row.cells[1].text = da.action
                row.cells[2].text = da.platform_or_method

        if roadmap_output.money_ask_message:
            _add_heading(doc, "Day 30 Money Ask", level=2)
            doc.add_paragraph(roadmap_output.money_ask_message)

        if roadmap_output.day30_success_definition:
            _add_heading(doc, "Day 30 — Success Definition", level=2)
            doc.add_paragraph(roadmap_output.day30_success_definition)


def export_build_guide(state: dict) -> dict:
    """
    LangGraph node — builds the complete build guide .docx from pipeline state.

    Reads:  state['idea_output'], state['env_profile'], state['service_bundle'],
            state['feature_bundle'], state['build_steps'], state['gtm_output'],
            state['roadmap_output'], state['blueprint_output'], state['project_slug']
    Writes: state['docx_path']
    """
    if not DOCX_AVAILABLE:
        safe_print("[DOCX EXPORTER] python-docx not available — skipping export.")
        state['docx_path'] = ""
        return state

    idea_output   = state.get('idea_output')
    env_profile   = state.get('env_profile')
    service_bundle = state.get('service_bundle')
    feature_bundle = state.get('feature_bundle')
    build_steps   = state.get('build_steps', [])
    gtm_output    = state.get('gtm_output')
    roadmap_output = state.get('roadmap_output')
    blueprint_output = state.get('blueprint_output')
    project_slug  = state.get('project_slug', '') or ''

    project_name = ""
    if idea_output and hasattr(idea_output, 'project_name'):
        project_name = idea_output.project_name or ''
    if not project_name:
        project_name = project_slug.replace("-", " ").title() or 'Forge Project'

    # Ensure we have a valid slug for file paths
    if not project_slug or project_slug.strip() == '':
        project_slug = project_name.lower().replace(' ', '-').replace('/', '-').replace('\\', '-')
    if not project_slug:
        project_slug = 'forge-project'

    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # SECTION 1: Cover Page
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(project_name.upper())
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = HEADING_FONT
    run.font.color.rgb = RGBColor(0x1E, 0x6F, 0xB5)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Complete Build Guide — Zero Ambiguity")
    sub_run.font.size = Pt(14)
    sub_run.font.name = BODY_FONT
    sub_run.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated by FORGE on {datetime.now().strftime('%d %B %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.name = BODY_FONT

    doc.add_paragraph()

    # Env profile summary
    if env_profile:
        _add_heading(doc, "Your Environment", level=2)
        profile_lines = [
            f"Operating System: {env_profile.os.title()}",
            f"AI Coding Tool: {env_profile.ai_cli.replace('_', ' ').title()}",
            f"Experience Level: {env_profile.experience.title()}",
            f"GPU: {env_profile.gpu}",
            f"Node.js installed: {'Yes' if env_profile.node_installed else 'No — install from nodejs.org'}",
            f"Python installed: {'Yes' if env_profile.python_installed else 'No — install from python.org'}",
        ]
        for line in profile_lines:
            doc.add_paragraph(line, style='List Bullet')
        doc.add_paragraph()

    # Idea summary
    if idea_output and not getattr(idea_output, 'failed', False):
        _add_heading(doc, "Project Summary", level=2)
        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = 'Light Grid'
        rows = [
            ("Job-to-be-Done", idea_output.job_to_be_done or ""),
            ("Target User", f"{idea_output.target_persona_name} — {idea_output.target_persona_description[:150]}"),
            ("Pain Score",  f"{idea_output.pain_score}/10"),
            ("Market Size", idea_output.market_size_estimate or ""),
            ("First Validation", idea_output.first_week_validation or ""),
        ]
        for row, (k, v) in zip(tbl.rows, rows):
            row.cells[0].paragraphs[0].add_run(k).bold = True
            row.cells[1].text = v
        doc.add_paragraph()

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # SECTION 2: Pre-Build Checklist
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_heading(doc, "Pre-Build Checklist", level=1)
    doc.add_paragraph(
        "Complete every item below BEFORE writing any code. "
        "This is the exact setup required for this project."
    )
    doc.add_paragraph()

    _build_service_table(doc, service_bundle)

    # ── .env template ────────────────────────────────────────────────────────
    _build_env_template(doc, service_bundle)

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # SECTION 3: Feature Build Guide (SETUP + CODING steps)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_heading(doc, "Build Guide", level=1)

    if blueprint_output and not getattr(blueprint_output, 'failed', False):
        _add_heading(doc, "MVP Definition", level=2)
        mvp_para = doc.add_paragraph()
        mvp_run = mvp_para.add_run(blueprint_output.mvp_definition or "")
        mvp_run.bold = True
        doc.add_paragraph()

        if blueprint_output.folder_structure:
            _add_heading(doc, "Folder Structure", level=2)
            _add_code_block(doc, blueprint_output.folder_structure)
            doc.add_paragraph()

    # ── Walk through build steps ──────────────────────────────────────────────
    if build_steps:
        _add_heading(doc, "Step-by-Step Build Instructions", level=2)
        setup_counter = 0
        coding_counter = 0

        for step in build_steps:
            step_type = step.step_type.value if hasattr(step.step_type, 'value') else str(step.step_type)

            if step_type == "setup":
                setup_counter += 1
                _add_step_badge(doc, "setup", step.label, setup_counter)

                for cmd in step.commands:
                    _add_code_block(doc, cmd)

                if step.files_created:
                    doc.add_paragraph(
                        f"Files created: {', '.join(step.files_created)}"
                    ).italic = True

            elif step_type == "coding":
                coding_counter += 1
                _add_step_badge(doc, "coding", step.label, coding_counter)

                # Time estimate
                doc.add_paragraph(f"⏱ Estimated time: {step.time_estimate_minutes} minutes")

                # AI prompt
                if step.ai_prompt:
                    _add_code_block(doc, step.ai_prompt)

                # Files to create
                if step.files_created:
                    doc.add_paragraph(
                        f"File to create: {', '.join(step.files_created)}"
                    ).italic = True

                # Env vars
                if step.env_vars_needed:
                    doc.add_paragraph(
                        f"Env vars: {', '.join(step.env_vars_needed)}"
                    ).italic = True

            # Success condition
            if step.success_condition:
                sc_para = doc.add_paragraph()
                sc_run = sc_para.add_run(f"✅ Success: {step.success_condition}")
                sc_run.bold = True
                sc_run.font.color.rgb = GREEN_BG

            doc.add_paragraph()
    else:
        # Fallback: render from blueprint tasks directly
        if blueprint_output and blueprint_output.mvp_tasks:
            _add_heading(doc, "Development Tasks", level=2)
            for task in blueprint_output.mvp_tasks:
                if task.tier != 'MVP':
                    continue
                _add_step_badge(doc, "coding", task.title, task.task_number)
                doc.add_paragraph(f"File: {task.file_path}")
                if task.gemini_cli_prompt:
                    _add_code_block(doc, task.gemini_cli_prompt)
                doc.add_paragraph(f"Test: {task.test_command}")
                sc_para = doc.add_paragraph()
                sc_para.add_run(f"✅ {task.success_check}").bold = True
                doc.add_paragraph()

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # SECTION 4: Deployment Guide
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _add_heading(doc, "Deployment Guide", level=1)

    os_name = env_profile.os if env_profile else "windows"
    _add_heading(doc, f"Deploy to Railway (Free Tier) — {os_name.title()}", level=2)

    deploy_steps_windows = [
        "npm install -g @railway/cli",
        "railway login",
        "railway init",
        "railway up",
        "railway open",
    ]
    deploy_steps_unix = [
        "npm install -g @railway/cli",
        "railway login",
        "railway init",
        "railway up",
        "railway open",
    ]
    deploy_steps = deploy_steps_windows if os_name == "windows" else deploy_steps_unix

    setup_num = (setup_counter if build_steps else 0) + 1
    _add_step_badge(doc, "setup", "Deploy to Railway", setup_num)
    for cmd in deploy_steps:
        _add_code_block(doc, cmd)

    sc_para = doc.add_paragraph()
    sc_para.add_run(
        "✅ Success: Railway dashboard shows 'Deployed' and public URL returns HTTP 200."
    ).bold = True
    doc.add_paragraph()

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # SECTION 5: GTM + 30-Day Roadmap
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _build_gtm_section(doc, gtm_output, roadmap_output)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # Save
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    output_dir = Path("outputs") / project_slug
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{project_slug}_build_guide.docx"

    doc.save(str(output_path))
    safe_print(f"[DOCX EXPORTER] [OK] Build guide saved: {output_path}")

    state['docx_path'] = str(output_path)
    return state